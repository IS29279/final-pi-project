"""
exports.py
Pi Intrusion Testing Appliance - Report Export Builders
Team 3 / ITP 258 Sprint 2

Generates downloadable versions of a scan report in four flavors:
  - Full report as PDF       (everything the web page shows)
  - Full report as Word doc  (same content, editable)
  - Executive summary PDF    (title, summary, "what to do next", top findings)
  - Executive summary Word   (same, editable)

Dependencies (pure Python, pip-installable, no system packages needed):
  pip install fpdf2 python-docx

Public functions:
  build_plain_english_guidance(flags)  → list of dicts describing actions in
                                          non-technical language

  build_full_pdf(context)              → bytes (PDF)
  build_full_docx(context)             → bytes (DOCX)
  build_summary_pdf(context)           → bytes (PDF)
  build_summary_docx(context)          → bytes (DOCX)

"context" for each builder is a dict with keys:
  scan, report, hosts_with_ports, traffic, flags, summary
(Same shape as the variables passed to report.html.)
"""

import io
import datetime
from collections import Counter


# ---------------------------------------------------------------------------
# Plain-English guidance builder
# ---------------------------------------------------------------------------
#
# Maps each flag type to a short, non-technical action statement. The table
# is keyword-based against the flag reason so it stays aligned with whatever
# reason strings FinalApp emits.
#
# Each entry is (keyword, kind, sentence) where kind is one of:
#   "action"       - concrete thing to do / ask IT to do
#   "reassurance"  - positive finding, no action needed
#
# We de-duplicate at the end — six RDP findings on six hosts should produce
# one action line, not six.
# ---------------------------------------------------------------------------

_PLAIN_ENGLISH_MAP = [
    # Critical / high actions
    ("telnet",           "action",      "Ask IT to turn off Telnet on any device where it's running. Telnet sends passwords in plain text, which means anyone on the network can read them."),
    ("smb",              "action",      "Ask IT to make sure file sharing (SMB) isn't open to the whole network. It's the most common way ransomware gets in."),
    ("rdp",              "action",      "Remote Desktop (RDP) should not be reachable from outside your office. Ask IT to put it behind a VPN or turn it off if no one needs it."),
    ("vnc",              "action",      "Screen-sharing (VNC) is running and should either be disabled or better secured. Ask IT to review it."),
    ("nfs",              "action",      "File-sharing over NFS is exposed. Ask IT to restrict who can connect to it."),
    ("ftp",              "action",      "FTP is running and sends passwords in plain text. Ask IT to use SFTP or HTTPS file transfer instead."),
    ("vsftpd",           "action",      "This FTP server has a known backdoor built into it. Ask IT to take it offline immediately and replace it."),
    ("outdated openssh", "action",      "The remote-login software (SSH) is running an outdated version with known security holes. Ask IT to update it."),
    ("outdated open",    "action",      "A critical encryption library (OpenSSL) is out of date and vulnerable. Ask IT to patch it."),
    ("end-of-life apache","action",     "The web server software is too old to receive security fixes anymore. Ask IT to upgrade or replace it."),
    ("end-of-life iis",  "action",      "The Windows web server software is outdated and has known serious vulnerabilities. Ask IT to upgrade the server."),
    ("iis 6",            "action",      "The Windows web server is running an ancient version that can be broken into remotely. This needs urgent attention."),
    ("mssql",            "action",      "A database (MSSQL) is accessible from the network. Databases should normally only be reachable by the applications that use them — ask IT to lock it down."),
    ("mysql",            "action",      "A database (MySQL) is accessible from the network. It should normally be hidden behind an application — ask IT to review."),
    ("postgres",         "action",      "A database (PostgreSQL) is accessible from the network. It should normally be hidden — ask IT to review."),
    ("default port 22",  "action",      "Remote-login (SSH) is running on its well-known default port, which makes it a target for automated attacks. Ask IT whether it should be moved or locked down to specific users."),
    ("http alternate",   "action",      "A web service is running on an unusual port. This is often a forgotten test server — ask IT to check what it is and whether it should still be running."),
    ("http (unencrypted","action",      "A website is running without encryption (HTTP, not HTTPS). Information sent to or from it can be intercepted. Ask IT to switch it to HTTPS."),

    # Traffic-level actions
    ("cleartext credentials observed", "action",
        "Someone on this network was seen typing a password into a system that didn't encrypt it. That password could have been read by anyone watching the network. Ask IT to find what system it was and switch it to encrypted authentication."),
    ("unencrypted traffic observed",    "action",
        "The scan saw old-style unencrypted traffic on the network. This isn't an immediate emergency, but modern traffic should all be encrypted — ask IT to look into which services are still plaintext."),
    ("only encrypted protocols observed", "reassurance",
        "The network scan only saw encrypted traffic during the check, which is what you want to see."),

    # Info tier reassurances
    ("https present",    "reassurance", "At least one device is using secure web traffic (HTTPS). That's the correct setup."),
    ("imaps present",    "reassurance", "Email retrieval is using an encrypted connection. That's the correct setup."),
    ("pop3s present",    "reassurance", "Email retrieval is using an encrypted connection. That's the correct setup."),
    ("ssh on non-standard port", "reassurance",
        "Someone took the time to move remote-login off its default port and keep the software current. That's a sign of a thoughtful setup."),
]


def build_plain_english_guidance(flags: list) -> list:
    """
    Turn the technical flag list into a de-duplicated list of plain-English
    action and reassurance statements a non-technical reader can understand.

    Returns a list of dicts:
        {"kind": "action" | "reassurance", "text": str, "severity": str}
    """
    seen_texts = set()
    actions = []
    reassurances = []

    # Walk flags in severity order so the "highest touched" severity on each
    # deduped action wins.
    severity_order = {"critical": 0, "high": 1, "medium": 2, "info": 3}
    sorted_flags = sorted(flags, key=lambda f: severity_order.get(f.get("severity", "info"), 99))

    for flag in sorted_flags:
        reason = (flag.get("reason") or "").lower()
        severity = flag.get("severity", "info")

        for keyword, kind, sentence in _PLAIN_ENGLISH_MAP:
            if keyword in reason:
                if sentence in seen_texts:
                    break
                seen_texts.add(sentence)
                entry = {"kind": kind, "text": sentence, "severity": severity}
                if kind == "action":
                    actions.append(entry)
                else:
                    reassurances.append(entry)
                break  # one plain-english entry per flag

    # Actions first (ordered by severity above), then reassurances
    return actions + reassurances


# ---------------------------------------------------------------------------
# PDF Builders (fpdf2)
# ---------------------------------------------------------------------------
#
# fpdf2 ships with Helvetica as a built-in Latin-1 font. Our data contains
# em-dashes and smart quotes (from reason strings like "Telnet — plaintext"),
# so we sanitize text to Latin-1-safe characters before writing. The visual
# result is nearly identical; em-dashes become " -- ", smart quotes become
# straight ones. Alternative would be shipping a Unicode TTF, but that adds
# a binary asset to the repo and complicates deployment.

_UNICODE_SANITIZE = {
    "\u2014": " -- ",  # em dash
    "\u2013": " - ",   # en dash
    "\u2018": "'",     # left single quote
    "\u2019": "'",     # right single quote / apostrophe
    "\u201C": '"',     # left double quote
    "\u201D": '"',     # right double quote
    "\u2026": "...",   # ellipsis
    "\u00A0": " ",     # non-breaking space
    "\u2022": "*",     # bullet
    "\u2192": "->",    # right arrow
    "\u2190": "<-",    # left arrow
    "\u00B7": "-",     # middle dot
}

def _sanitize(text):
    """Replace non-Latin-1 punctuation with ASCII equivalents for fpdf2."""
    if text is None:
        return ""
    text = str(text)
    for u, replacement in _UNICODE_SANITIZE.items():
        text = text.replace(u, replacement)
    # Fall back: strip any remaining non-Latin-1 characters so fpdf doesn't crash
    return text.encode("latin-1", "replace").decode("latin-1")


def _pdf_header(pdf, title, session_id):
    """Shared header block for PDF exports."""
    pdf.set_fill_color(88, 166, 255)  # Network Lock blue
    pdf.rect(0, 0, 210, 18, "F")
    pdf.set_xy(10, 4)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 5, _sanitize("NETWORK LOCK SECURITY"), ln=True)
    pdf.set_x(10)
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 5, _sanitize(title), ln=True)
    pdf.set_text_color(50, 50, 50)
    pdf.set_y(22)
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 4, _sanitize(f"Session {session_id[:8]} -- Generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"), ln=True)
    pdf.ln(3)
    pdf.set_text_color(30, 30, 30)


def _pdf_section(pdf, title):
    """Horizontal rule + section heading."""
    pdf.ln(2)
    pdf.set_draw_color(88, 166, 255)
    pdf.set_line_width(0.5)
    y = pdf.get_y()
    pdf.line(10, y, 200, y)
    pdf.ln(2)
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_text_color(88, 100, 120)
    pdf.cell(0, 6, _sanitize(title.upper()), ln=True)
    pdf.set_text_color(30, 30, 30)
    pdf.ln(1)


_SEVERITY_RGB = {
    "critical": (207, 34, 46),    # red
    "high":     (188, 76, 0),     # orange
    "medium":   (154, 103, 0),    # yellow-brown
    "info":     (9, 105, 218),    # blue
}


def _pdf_severity_pill(pdf, text, severity, w=22):
    """Draw a small colored severity pill at the current cursor."""
    r, g, b = _SEVERITY_RGB.get(severity, (100, 100, 100))
    pdf.set_fill_color(r, g, b)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 7)
    pdf.cell(w, 4, _sanitize(text.upper()), border=0, fill=True, align="C")
    pdf.set_text_color(30, 30, 30)


def _pdf_flag_table(pdf, flags, include_search_terms=True):
    """Render the flagged-concerns list grouped by severity."""
    severity_order = ["critical", "high", "medium", "info"]
    pdf.set_font("Helvetica", "", 9)

    for sev in severity_order:
        group = [f for f in flags if f["severity"] == sev]
        if not group:
            continue

        # Section header for this severity
        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 9)
        r, g, b = _SEVERITY_RGB[sev]
        pdf.set_text_color(r, g, b)
        pdf.cell(0, 5, _sanitize(f"{sev.upper()}  ({len(group)} finding{'s' if len(group) != 1 else ''})"), ln=True)
        pdf.set_text_color(30, 30, 30)

        for flag in group:
            target = flag["host"]
            if flag.get("port") is not None:
                target = f"{flag['host']}:{flag['port']}"

            # Line 1: target (bold)
            pdf.set_x(12)
            pdf.set_font("Helvetica", "B", 9)
            pdf.cell(0, 4.5, _sanitize(target), ln=True)

            # Line 2+: reason (wrapped, full width)
            pdf.set_x(14)
            pdf.set_font("Helvetica", "", 9)
            # Use a width of 180mm for the reason so it wraps before the right margin
            pdf.multi_cell(w=185, h=4.5, text=_sanitize(flag["reason"]))

            # Line 3+: search terms (indented, italicized, smaller)
            if include_search_terms and flag.get("search_terms"):
                pdf.set_x(14)
                pdf.set_font("Helvetica", "I", 8)
                pdf.set_text_color(100, 100, 100)
                terms = " | ".join(flag["search_terms"])
                pdf.multi_cell(w=185, h=4, text=_sanitize(f"Search: {terms}"))
                pdf.set_text_color(30, 30, 30)

            pdf.ln(1)


def _pdf_guidance(pdf, guidance):
    """Render the plain-English 'What to do next' guidance."""
    actions      = [g for g in guidance if g["kind"] == "action"]
    reassurances = [g for g in guidance if g["kind"] == "reassurance"]

    if actions:
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(207, 34, 46)
        pdf.cell(0, 5, _sanitize("What to do next"), ln=True)
        pdf.set_text_color(30, 30, 30)
        pdf.set_font("Helvetica", "", 9)
        for item in actions:
            pdf.set_x(12)
            pdf.multi_cell(0, 4.5, _sanitize(f"* {item['text']}"))
            pdf.ln(0.5)

    if reassurances:
        pdf.ln(2)
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(9, 105, 218)
        pdf.cell(0, 5, _sanitize("What's working well"), ln=True)
        pdf.set_text_color(30, 30, 30)
        pdf.set_font("Helvetica", "", 9)
        for item in reassurances:
            pdf.set_x(12)
            pdf.multi_cell(0, 4.5, _sanitize(f"* {item['text']}"))
            pdf.ln(0.5)


def build_full_pdf(ctx) -> bytes:
    """Full report as PDF -- everything the web page shows."""
    from fpdf import FPDF

    scan     = ctx["scan"]
    report   = ctx["report"]
    flags    = ctx["flags"]
    summary  = ctx["summary"]
    hosts    = ctx["hosts_with_ports"]
    traffic  = ctx["traffic"]
    guidance = build_plain_english_guidance(flags)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    _pdf_header(pdf, "Network Security Assessment Report -- Full", report["session_id"])

    # ── Scan metadata ──
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 4, _sanitize(f"Target: {scan['target_cidr']}"), ln=True)
    pdf.cell(0, 4, _sanitize(f"Status: {scan['status']}   |   Hosts scanned: {len(hosts)}   |   Total findings: {len(flags)}"), ln=True)
    pdf.ln(2)

    # ── Summary ──
    _pdf_section(pdf, "Summary of Findings")
    pdf.set_font("Helvetica", "B", 11)
    pdf.multi_cell(0, 5, _sanitize(summary["headline"]))
    pdf.ln(1)
    pdf.set_font("Helvetica", "", 9)
    pdf.multi_cell(0, 5, _sanitize(summary["paragraph"]))

    # ── Plain-English guidance ──
    if guidance:
        _pdf_section(pdf, "Here's What You Should Know")
        _pdf_guidance(pdf, guidance)

    # ── Flagged Concerns ──
    _pdf_section(pdf, "Flagged Concerns")
    if flags:
        _pdf_flag_table(pdf, flags, include_search_terms=True)
    else:
        pdf.set_font("Helvetica", "I", 9)
        pdf.cell(0, 5, _sanitize("No concerns flagged for this scan."), ln=True)

    # ── Discovered Hosts ──
    _pdf_section(pdf, "Discovered Hosts")
    if hosts:
        for item in hosts:
            h = item["host"]
            pdf.set_font("Helvetica", "B", 10)
            line = h["ip_address"]
            if h.get("hostname"):
                line += f"   ({h['hostname']})"
            pdf.cell(0, 5, _sanitize(line), ln=True)
            pdf.set_font("Helvetica", "", 8)
            if item["ports"]:
                for port in item["ports"]:
                    svc = port.get("service_name") or "-"
                    ver = port.get("service_version") or ""
                    pdf.set_x(14)
                    pdf.cell(0, 4, _sanitize(f"{port['port_number']}/{port['protocol']}  {svc}  {ver}"), ln=True)
            else:
                pdf.set_x(14)
                pdf.set_font("Helvetica", "I", 8)
                pdf.cell(0, 4, _sanitize("(no open ports detected)"), ln=True)
            pdf.ln(1)
    else:
        pdf.set_font("Helvetica", "I", 9)
        pdf.cell(0, 5, _sanitize("No hosts discovered."), ln=True)

    # ── Traffic Analysis ──
    _pdf_section(pdf, "Traffic Analysis")
    if traffic:
        pdf.set_font("Helvetica", "", 9)
        for t in traffic:
            creds = "YES -- review immediately" if t.get("cleartext_creds_found") else "None detected"
            pdf.cell(0, 4, _sanitize(f"PCAP: {t.get('pcap_path', '-')}"), ln=True)
            pdf.cell(0, 4, _sanitize(f"Protocols: {t.get('protocol_summary') or '-'}"), ln=True)
            pdf.cell(0, 4, _sanitize(f"Cleartext credentials: {creds}"), ln=True)
            pdf.ln(1)
    else:
        pdf.set_font("Helvetica", "I", 9)
        pdf.cell(0, 5, _sanitize("No traffic findings recorded."), ln=True)

    # fpdf2 returns bytearray; make it bytes for Flask
    return bytes(pdf.output())


def build_summary_pdf(ctx) -> bytes:
    """Executive-summary PDF: title, summary paragraph, guidance, top findings."""
    from fpdf import FPDF

    scan     = ctx["scan"]
    report   = ctx["report"]
    flags    = ctx["flags"]
    summary  = ctx["summary"]
    hosts    = ctx["hosts_with_ports"]
    guidance = build_plain_english_guidance(flags)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    _pdf_header(pdf, "Network Security Assessment -- Executive Summary", report["session_id"])

    # ── Scan context ──
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 4, _sanitize(f"Target: {scan['target_cidr']}   |   Hosts scanned: {len(hosts)}   |   Total findings: {len(flags)}"), ln=True)
    pdf.ln(2)

    # ── Summary ──
    _pdf_section(pdf, "Summary of Findings")
    pdf.set_font("Helvetica", "B", 12)
    pdf.multi_cell(0, 6, _sanitize(summary["headline"]))
    pdf.ln(1)
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 5, _sanitize(summary["paragraph"]))

    # ── What to do next ──
    if guidance:
        _pdf_section(pdf, "Here's What You Should Know")
        _pdf_guidance(pdf, guidance)

    # ── Top findings only (critical + high, cap at 10) ──
    top_flags = [f for f in flags if f["severity"] in ("critical", "high")][:10]
    if top_flags:
        _pdf_section(pdf, "Top Findings")
        _pdf_flag_table(pdf, top_flags, include_search_terms=False)

    return bytes(pdf.output())


# ---------------------------------------------------------------------------
# DOCX Builders (python-docx)
# ---------------------------------------------------------------------------

def _docx_add_header(doc, title, session_id):
    """Title + session identifier block for a docx."""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Top brand bar as a one-cell shaded paragraph
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("NETWORK LOCK SECURITY")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x58, 0xA6, 0xFF)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(title)
    sub_run.font.size = Pt(11)
    sub_run.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"Session {session_id[:8]} — Generated {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _docx_section(doc, title):
    """Add a section heading with consistent styling."""
    from docx.shared import Pt, RGBColor
    doc.add_paragraph()  # spacer
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x58, 0x64, 0x78)


def _docx_severity_label(severity):
    """Return short severity tag for prefixing list items."""
    return f"[{severity.upper()}]"


def _docx_add_flags(doc, flags, include_search_terms=True):
    """Write flagged concerns grouped by severity."""
    from docx.shared import Pt, RGBColor

    severity_order = ["critical", "high", "medium", "info"]
    sev_colors = {
        "critical": RGBColor(0xCF, 0x22, 0x2E),
        "high":     RGBColor(0xBC, 0x4C, 0x00),
        "medium":   RGBColor(0x9A, 0x67, 0x00),
        "info":     RGBColor(0x09, 0x69, 0xDA),
    }

    for sev in severity_order:
        group = [f for f in flags if f["severity"] == sev]
        if not group:
            continue

        # Sev group heading
        p = doc.add_paragraph()
        run = p.add_run(f"{sev.upper()}  ({len(group)} finding{'s' if len(group) != 1 else ''})")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = sev_colors[sev]

        for flag in group:
            target = flag["host"]
            if flag.get("port") is not None:
                target = f"{flag['host']}:{flag['port']}"

            p = doc.add_paragraph(style="List Bullet")
            target_run = p.add_run(f"{target}  ")
            target_run.bold = True
            p.add_run(flag["reason"])

            if include_search_terms and flag.get("search_terms"):
                sp = doc.add_paragraph()
                sp.paragraph_format.left_indent = Pt(18)
                sr = sp.add_run(f"Search: {' | '.join(flag['search_terms'])}")
                sr.font.size = Pt(9)
                sr.italic = True
                sr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def _docx_add_guidance(doc, guidance):
    from docx.shared import Pt, RGBColor
    actions      = [g for g in guidance if g["kind"] == "action"]
    reassurances = [g for g in guidance if g["kind"] == "reassurance"]

    if actions:
        p = doc.add_paragraph()
        run = p.add_run("What to do next")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xCF, 0x22, 0x2E)
        for item in actions:
            doc.add_paragraph(item["text"], style="List Bullet")

    if reassurances:
        p = doc.add_paragraph()
        run = p.add_run("What's working well")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x09, 0x69, 0xDA)
        for item in reassurances:
            doc.add_paragraph(item["text"], style="List Bullet")


def build_full_docx(ctx) -> bytes:
    """Full report as .docx — everything the web page shows."""
    from docx import Document
    from docx.shared import Pt

    scan     = ctx["scan"]
    report   = ctx["report"]
    flags    = ctx["flags"]
    summary  = ctx["summary"]
    hosts    = ctx["hosts_with_ports"]
    traffic  = ctx["traffic"]
    guidance = build_plain_english_guidance(flags)

    doc = Document()
    _docx_add_header(doc, "Network Security Assessment Report — Full", report["session_id"])

    # Scan metadata
    p = doc.add_paragraph()
    p.add_run(f"Target: ").bold = True
    p.add_run(f"{scan['target_cidr']}    ")
    p.add_run("Status: ").bold = True
    p.add_run(f"{scan['status']}    ")
    p.add_run("Hosts: ").bold = True
    p.add_run(f"{len(hosts)}    ")
    p.add_run("Findings: ").bold = True
    p.add_run(f"{len(flags)}")

    # Summary
    _docx_section(doc, "Summary of Findings")
    hp = doc.add_paragraph()
    hr = hp.add_run(summary["headline"])
    hr.bold = True
    hr.font.size = Pt(13)
    doc.add_paragraph(summary["paragraph"])

    # Guidance
    if guidance:
        _docx_section(doc, "Here's What You Should Know")
        _docx_add_guidance(doc, guidance)

    # Flagged Concerns
    _docx_section(doc, "Flagged Concerns")
    if flags:
        _docx_add_flags(doc, flags, include_search_terms=True)
    else:
        doc.add_paragraph("No concerns flagged for this scan.")

    # Discovered Hosts
    _docx_section(doc, "Discovered Hosts")
    if hosts:
        for item in hosts:
            h = item["host"]
            line = h["ip_address"]
            if h.get("hostname"):
                line += f"   ({h['hostname']})"
            p = doc.add_paragraph()
            p.add_run(line).bold = True
            if item["ports"]:
                for port in item["ports"]:
                    svc = port.get("service_name") or "-"
                    ver = port.get("service_version") or ""
                    doc.add_paragraph(
                        f"{port['port_number']}/{port['protocol']}  {svc}  {ver}",
                        style="List Bullet"
                    )
            else:
                doc.add_paragraph("(no open ports detected)", style="List Bullet")
    else:
        doc.add_paragraph("No hosts discovered.")

    # Traffic Analysis
    _docx_section(doc, "Traffic Analysis")
    if traffic:
        for t in traffic:
            creds = "YES — review immediately" if t.get("cleartext_creds_found") else "None detected"
            doc.add_paragraph(f"PCAP: {t.get('pcap_path', '-')}", style="List Bullet")
            doc.add_paragraph(f"Protocols: {t.get('protocol_summary') or '-'}", style="List Bullet")
            doc.add_paragraph(f"Cleartext credentials: {creds}", style="List Bullet")
    else:
        doc.add_paragraph("No traffic findings recorded.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_summary_docx(ctx) -> bytes:
    """Executive-summary .docx."""
    from docx import Document
    from docx.shared import Pt

    scan     = ctx["scan"]
    report   = ctx["report"]
    flags    = ctx["flags"]
    summary  = ctx["summary"]
    hosts    = ctx["hosts_with_ports"]
    guidance = build_plain_english_guidance(flags)

    doc = Document()
    _docx_add_header(doc, "Network Security Assessment — Executive Summary", report["session_id"])

    # Scan metadata
    p = doc.add_paragraph()
    p.add_run(f"Target: ").bold = True
    p.add_run(f"{scan['target_cidr']}    ")
    p.add_run("Hosts: ").bold = True
    p.add_run(f"{len(hosts)}    ")
    p.add_run("Findings: ").bold = True
    p.add_run(f"{len(flags)}")

    # Summary
    _docx_section(doc, "Summary of Findings")
    hp = doc.add_paragraph()
    hr = hp.add_run(summary["headline"])
    hr.bold = True
    hr.font.size = Pt(13)
    doc.add_paragraph(summary["paragraph"])

    # Guidance
    if guidance:
        _docx_section(doc, "Here's What You Should Know")
        _docx_add_guidance(doc, guidance)

    # Top findings only
    top_flags = [f for f in flags if f["severity"] in ("critical", "high")][:10]
    if top_flags:
        _docx_section(doc, "Top Findings")
        _docx_add_flags(doc, top_flags, include_search_terms=False)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
